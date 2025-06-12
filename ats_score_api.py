from flask import Flask, request, jsonify
import pdfplumber
from docx import Document
import spacy
import re
import os
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity
import requests
import uuid
import tempfile

app = Flask(__name__)

# Load spaCy model
nlp = spacy.load("en_core_web_md")

# Fixed weights for consistent scoring
WEIGHTS = {"keyword": 0.5, "semantic": 0.4, "context": 0.1}

def download_file(url, temp_dir):
    """Download file from URL and save to temp directory."""
    try:
        response = requests.get(url, stream=True)
        if response.status_code != 200:
            return None, f"Failed to download file from {url}: Status {response.status_code}"
        
        # Determine file extension from URL
        if url.lower().endswith('.pdf'):
            ext = '.pdf'
        elif url.lower().endswith('.docx'):
            ext = '.docx'
        else:
            return None, "URL must point to a PDF or Word (.docx) file"
        
        # Save to temp file
        temp_path = os.path.join(temp_dir, f"file_{uuid.uuid4()}{ext}")
        with open(temp_path, 'wb') as f:
            for chunk in response.iter_content(chunk_size=8192):
                if chunk:
                    f.write(chunk)
        return temp_path, None
    except Exception as e:
        return None, str(e)

def extract_text_from_pdf(file_path):
    """Extract text from a PDF file."""
    try:
        with pdfplumber.open(file_path) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text
            if not text.strip():
                raise ValueError("No extractable text found")
            return text.strip(), None
    except Exception as e:
        return None, str(e)

def extract_text_from_docx(file_path):
    """Extract text from a Word (.docx) file."""
    try:
        doc = Document(file_path)
        text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        for section in doc.sections:
            for header in section.header.paragraphs:
                if header.text.strip():
                    text += header.text + "\n"
            for footer in section.footer.paragraphs:
                if footer.text.strip():
                    text += footer.text + "\n"
        if not text.strip():
            raise ValueError("No extractable text found")
        return text.strip(), None
    except Exception as e:
        return None, str(e)

def preprocess_text(text):
    """Preprocess text using spaCy."""
    if not text:
        return []
    text = text.lower()
    text = re.sub(r'[^\w\s]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    doc = nlp(text)
    tokens = [token.lemma_ for token in doc if token.pos_ in ('NOUN', 'VERB', 'PROPN') and not token.is_stop and len(token.text) > 2]
    return tokens

def get_context_similarity(resume_text, jd_text):
    """Calculate sentence-level context similarity."""
    resume_doc = nlp(resume_text)
    jd_doc = nlp(jd_text)
    sentences_resume = [sent.text for sent in resume_doc.sents]
    sentences_jd = [sent.text for sent in jd_doc.sents]
    if not sentences_resume or not sentences_jd:
        return 0
    similarities = []
    for sent_resume in sentences_resume[:10]:
        sent_resume_doc = nlp(sent_resume)
        for sent_jd in sentences_jd[:10]:
            sent_jd_doc = nlp(sent_jd)
            sim = sent_resume_doc.similarity(sent_jd_doc)
            if sim > 0:
                similarities.append(sim)
    return np.mean(similarities) * 100 if similarities else 0

@app.route('/calculate-ats', methods=['POST'])
def calculate_ats_score():
    """Calculate ATS score from resume and JD URLs."""
    # Check for JSON payload
    if not request.is_json:
        return jsonify({"error": "Request must be JSON"}), 400
    
    data = request.get_json()
    jd_url = data.get("Job Description URL")
    resume_url = data.get("Resume URL")
    applicant_id = data.get("applicant_id")

    # Validate input
    if not all([jd_url, resume_url, applicant_id]):
        return jsonify({"error": "Missing required fields: Job Description URL, Resume URL, applicant_id"}), 400

    # Create temporary directory
    with tempfile.TemporaryDirectory() as temp_dir:
        # Download files
        resume_path, resume_error = download_file(resume_url, temp_dir)
        if resume_error:
            return jsonify({"error": resume_error}), 400
        
        jd_path, jd_error = download_file(jd_url, temp_dir)
        if jd_error:
            return jsonify({"error": jd_error}), 400

        # Extract text
        resume_text, resume_error = (extract_text_from_pdf(resume_path) if resume_path.endswith('.pdf')
                                    else extract_text_from_docx(resume_path))
        jd_text, jd_error = (extract_text_from_pdf(jd_path) if jd_path.endswith('.pdf')
                             else extract_text_from_docx(jd_path))

        if resume_error or jd_error:
            return jsonify({"error": resume_error or jd_error}), 400
        if not resume_text or not jd_text:
            return jsonify({"error": "Failed to extract text from one or both files"}), 400

        # Preprocess texts
        resume_tokens = preprocess_text(resume_text)
        jd_tokens = preprocess_text(jd_text)

        if not resume_tokens or not jd_tokens:
            return jsonify({"error": "No valid content found in one or both files"}), 400

        # Keyword Matching
        jd_keywords = set(jd_tokens)
        matched_keywords = [token for token in resume_tokens if token in jd_keywords]
        keyword_score = (len(matched_keywords) / len(jd_keywords)) * 100 if jd_keywords else 0

        # Semantic Similarity
        resume_doc = nlp(resume_text)
        jd_doc = nlp(jd_text)
        semantic_score = resume_doc.similarity(jd_doc) * 100

        # Context Similarity
        context_score = get_context_similarity(resume_text, jd_text)

        # Combined ATS Score
        ats_score = (
            WEIGHTS["keyword"] * keyword_score +
            WEIGHTS["semantic"] * semantic_score +
            WEIGHTS["context"] * context_score
        )

        # Round score
        ats_score = round(ats_score, 2)

        # Response
        response = {
            "applicant_id": applicant_id,
            "ats_score": ats_score
        }

        return jsonify(response), 200

if __name__ == "__main__":
    app.run(debug=True, host="0.0.0.0", port=5000)